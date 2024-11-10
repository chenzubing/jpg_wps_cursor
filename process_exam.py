from docx import Document
from docx.shared import RGBColor
import re
from difflib import SequenceMatcher
import openai  # 需要添加OpenAI支持
from dotenv import load_dotenv
import os

load_dotenv()

def similarity_ratio(str1, str2):
    """计算两个字符串的相似度"""
    return SequenceMatcher(None, str1, str2).ratio()

def remove_duplicates(questions, similarity_threshold=0.85):
    """去除重复的题目"""
    unique_questions = []
    for q in questions:
        is_duplicate = False
        for existing_q in unique_questions:
            if similarity_ratio(q['question'], existing_q['question']) > similarity_threshold:
                is_duplicate = True
                break
        if not is_duplicate:
            unique_questions.append(q)
    return unique_questions

def process_with_llm(questions):
    """使用LLM处理题目内容"""
    processed_questions = []
    
    # 从环境变量获取API密钥
    openai.api_key = os.getenv('OPENAI_API_KEY')
    
    if not openai.api_key:
        raise ValueError("请在.env文件中设置OPENAI_API_KEY")
    
    for q in questions:
        # 构建提示
        prompt = f"""
        请分析以下试题，并提供规范化的格式：
        
        题目：{q['question']}
        选项：
        {chr(10).join([opt['text'] for opt in q['options']])}
        正确答案：{', '.join(q['correct_answers'])}
        
        请返回：
        1. 规范化的题目描述
        2. 确认正确答案是否合理
        3. 如果发现题目或选项有问题，请指出
        """
        
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "你是一个专业的考试题目审核助手。"},
                    {"role": "user", "content": prompt}
                ]
            )
            
            # 处理LLM的响应
            llm_response = response.choices[0].message.content
            
            # 将LLM的分析结果添加到题目数据中
            q['llm_analysis'] = llm_response
            processed_questions.append(q)
            
        except Exception as e:
            print(f"处理题目时出错: {str(e)}")
            processed_questions.append(q)
    
    return processed_questions

def extract_questions_from_docx(file_path):
    doc = Document(file_path)
    questions = []
    current_question = None
    current_options = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
            
        # 检查是否是新的题目（通常以数字开头）
        if re.match(r'^\d+[\.)、]', text):
            # 如果已有题目，保存之前的题目
            if current_question:
                questions.append({
                    'question': current_question,
                    'options': current_options,
                    'correct_answers': get_correct_answers(current_options)
                })
            
            current_question = text
            current_options = []
        # 检查是否是选项（通常以A、B、C、D开头）
        elif re.match(r'^[A-Z][\.)、]', text):
            # 检查这个选项是否被标记为彩色
            is_correct = False
            for run in paragraph.runs:
                if run.font.color.rgb:  # 如果文字有颜色
                    is_correct = True
                    break
            
            current_options.append({
                'text': text,
                'is_correct': is_correct
            })
    
    # 添加最后一个题目
    if current_question:
        questions.append({
            'question': current_question,
            'options': current_options,
            'correct_answers': get_correct_answers(current_options)
        })
    
    return questions

def get_correct_answers(options):
    return [opt['text'][0] for opt in options if opt['is_correct']]

def format_output(questions):
    doc = Document()
    for i, q in enumerate(questions, 1):
        # 添加题目
        doc.add_paragraph(f"{i}. {q['question']}")
        
        # 添加选项
        for option in q['options']:
            p = doc.add_paragraph(option['text'])
            if option['is_correct']:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 0, 0)
        
        # 添加正确答案说明
        answers = '、'.join(q['correct_answers'])
        doc.add_paragraph(f"正确答案：{answers}")
        
        # 添加LLM分析结果（如果有）
        if 'llm_analysis' in q:
            doc.add_paragraph("AI分析：").bold = True
            doc.add_paragraph(q['llm_analysis'])
        
        doc.add_paragraph()  # 分隔行

def main():
    input_file = 'temp.docx'
    output_file = 'processed_exam.docx'
    
    # 1. 提取题目
    questions = extract_questions_from_docx(input_file)
    
    # 2. 去除重复题目
    unique_questions = remove_duplicates(questions)
    print(f"去除重复后的题目数量: {len(unique_questions)}")
    
    # 3. 使用LLM处理题目
    processed_questions = process_with_llm(unique_questions)
    
    # 4. 生成输出文档
    output_doc = format_output(processed_questions)
    output_doc.save(output_file)
    print(f"处理完成！结果已保存到 {output_file}")

if __name__ == '__main__':
    main() 